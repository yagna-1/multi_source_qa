# ---------------------------
# app.py (Main Application)
# ---------------------------
"""
Multi-Source Document QA System with Enhanced Embedding Fine-Tuning and Large Document Handling

This Streamlit application processes multiple document types (PDF, DOCX, TXT),
splits the text into chunks, indexes them using FAISS with embeddings from a cached
SentenceTransformer, and then uses a LangChain-powered QA chain to answer queries.
It supports fine-tuning of the embedding model on domain-specific training data and
has special handling for large documents (over 1MB).

Features:
- Multi-format document processing with fallback and large-document handling.
- Cached SentenceTransformer embedder.
- FAISS indexing and custom retriever integrated with LangChain.
- Asynchronous QA chain initialization with exponential backoff.
- Evaluation metrics (fuzzy matching-based Recall@k and BLEU).
- Embedding fine-tuning via a training dataset.
- Session state caching for persistence across interactions.
- Detailed logging and error handling.
"""

import os
import time
import logging
import tempfile
import traceback
import numpy as np
import faiss
import streamlit as st
import asyncio
from typing import List, Any, Tuple
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.llms import OpenAI
from langchain.chains import RetrievalQA
from sentence_transformers import SentenceTransformer
from PyPDF2 import PdfReader
from pdfplumber import PDF
import docx
import difflib  # For fuzzy matching as an additional evaluation metric
from nltk.translate.bleu_score import sentence_bleu

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------------------
# Configuration
# ---------------------------
CONFIG = {
    "chunk_size": int(os.getenv("CHUNK_SIZE", 500)),
    "chunk_overlap": int(os.getenv("CHUNK_OVERLAP", 50)),
    "max_file_size_mb": int(os.getenv("MAX_FILE_SIZE_MB", 20)),
    "retrieval_k": int(os.getenv("RETRIEVAL_K", 3)),
    "faiss_index_type": os.getenv("FAISS_INDEX_TYPE", "L2")  # Placeholder for future index types
}

# ---------------------------
# Cached Resource: Sentence Transformer Model
# ---------------------------
@st.cache_resource(show_spinner=False)
def get_embedder() -> SentenceTransformer:
    """Load and cache the SentenceTransformer model."""
    return SentenceTransformer('all-MiniLM-L6-v2')

# ---------------------------
# Caching: Heavy Operations
# ---------------------------
@st.cache_data(show_spinner=False)
def cached_extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Cache the extraction of text from a file to avoid redundant processing.
    """
    from io import BytesIO
    file_obj = BytesIO(file_bytes)
    file_obj.type = file_type  # Simulate file MIME type
    processor = DocumentProcessor()
    return processor.extract_text(file_obj)

# ---------------------------
# Document Processing Class
# ---------------------------
class DocumentProcessor:
    """
    A class for extracting text from various file formats and splitting it into chunks.
    """
    def __init__(self, chunk_size: int = CONFIG["chunk_size"], chunk_overlap: int = CONFIG["chunk_overlap"]):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len
        )
        
    def extract_text(self, file: Any) -> str:
        """
        Extract text from a file-like object based on its MIME type.
        
        Args:
            file (Any): A file-like object with a 'type' attribute.
        
        Returns:
            str: The extracted text.
        """
        try:
            if file.type == "application/pdf":
                return self._extract_from_pdf(file)
            elif file.type == "text/plain":
                return file.read().decode("utf-8", errors="ignore")
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_from_docx(file)
            else:
                raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")
        except Exception as e:
            logger.error(f"Error processing file: {str(e)}\n{traceback.format_exc()}")
            raise RuntimeError("Error processing document. Please check the file format and try again.")

    def _extract_from_pdf(self, file: Any) -> str:
        """
        Extract text from a PDF file, using a fallback to pdfplumber if needed.
        """
        try:
            reader = PdfReader(file)
            text = "".join([page.extract_text() or "" for page in reader.pages])
            if not text.strip():
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(file.getbuffer())
                    tmp.flush()
                    with PDF.open(tmp.name) as pdf:
                        text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                os.unlink(tmp.name)
            return text
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}\n{traceback.format_exc()}")
            raise RuntimeError("Failed to extract text from PDF.")

    def _extract_from_docx(self, file: Any) -> str:
        """
        Extract text from a DOCX file.
        """
        try:
            doc = docx.Document(file)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}\n{traceback.format_exc()}")
            raise RuntimeError("Failed to extract text from DOCX file.")

    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks after verifying it is not empty.
        
        Args:
            text (str): The text to split.
        
        Returns:
            List[str]: List of text chunks.
        """
        if not text.strip():
            raise ValueError("The document contains no text.")
        return self.text_splitter.split_text(text)

    def _handle_large_document(self, text: str) -> List[str]:
        """
        Process documents larger than 1MB using a LangChain TextLoader.
        
        Args:
            text (str): The document text.
            
        Returns:
            List[str]: A list of text chunks from the large document.
        """
        from langchain.document_loaders import TextLoader
        with tempfile.NamedTemporaryFile(delete=False, mode="w", encoding="utf-8") as tmp:
            tmp.write(text)
            tmp_path = tmp.name
        try:
            loader = TextLoader(tmp_path)
            # load_and_split uses the provided text splitter to split the document
            return loader.load_and_split(self.text_splitter)
        finally:
            os.unlink(tmp_path)

# ---------------------------
# Question Answering System Class
# ---------------------------
class QASystem:
    """
    A QA system that builds a FAISS index from text chunks, creates a custom retriever,
    integrates with a LangChain QA chain, and provides evaluation metrics.
    """
    def __init__(self):
        self.embedder = get_embedder()  # Cached embedder
        self.index = None
        self.chunks: List[str] = []
        self.qa_chain = None
        
    def build_index(self, chunks: List[str]) -> None:
        """
        Create a FAISS index from the provided text chunks.
        
        Args:
            chunks (List[str]): List of text chunks.
        """
        if not chunks:
            raise ValueError("No text chunks provided for indexing.")
            
        embeddings = self.embedder.encode(chunks)
        dimension = embeddings.shape[1]
        if CONFIG["faiss_index_type"] == "L2":
            self.index = faiss.IndexFlatL2(dimension)
        else:
            self.index = faiss.IndexFlatL2(dimension)
        self.index.add(np.array(embeddings).astype('float32'))
        self.chunks = chunks
        
    async def initialize_qa_chain(self) -> None:
        """
        Initialize the LangChain QA system asynchronously with retry logic.
        """
        max_retries = 3
        for attempt in range(max_retries):
            try:
                self.qa_chain = RetrievalQA.from_chain_type(
                    llm=OpenAI(model_name="gpt-3.5-turbo", temperature=0),
                    chain_type="stuff",
                    retriever=self._create_custom_retriever()
                )
                return
            except Exception as e:
                logger.error(f"QA chain initialization failed on attempt {attempt + 1}: {str(e)}\n{traceback.format_exc()}")
                if attempt == max_retries - 1:
                    raise RuntimeError("Failed to initialize QA system after multiple attempts. Please try again later.")
                await asyncio.sleep(2 ** attempt)
                
    def _create_custom_retriever(self) -> Any:
        """
        Create a custom FAISS-based retriever that integrates with LangChain.
        
        Returns:
            Any: A custom retriever instance.
        """
        from langchain.schema import BaseRetriever, Document
        
        class FAISSRetriever(BaseRetriever):
            def __init__(self, index: Any, chunks: List[str], embedder: SentenceTransformer):
                self.index = index
                self.chunks = chunks
                self.embedder = embedder
                
            def get_relevant_documents(self, query: str) -> List[Any]:
                query_embedding = self.embedder.encode([query])
                distances, indices = self.index.search(np.array(query_embedding).astype('float32'), k=CONFIG["retrieval_k"])
                docs = []
                for idx in indices[0]:
                    if idx < len(self.chunks):
                        docs.append(Document(page_content=self.chunks[idx]))
                return docs
                
        return FAISSRetriever(self.index, self.chunks, self.embedder)

    def fine_tune_embeddings(self, training_data: List[Tuple[str, str]]) -> None:
        """
        Fine-tune the SentenceTransformer embedding model on domain-specific training data.
        
        Args:
            training_data (List[Tuple[str, str]]): A list of question-answer pairs.
        """
        from sentence_transformers import InputExample, losses
        from torch.utils.data import DataLoader
        
        examples = [InputExample(texts=[q, a]) for q, a in training_data]
        train_dataloader = DataLoader(examples, batch_size=16)
        train_loss = losses.MultipleNegativesRankingLoss(self.embedder)
        
        self.embedder.fit(
            train_objectives=[(train_dataloader, train_loss)],
            epochs=3,
            show_progress_bar=True
        )

    def _fuzzy_match(self, a: str, b: str) -> float:
        """Compute fuzzy similarity between two strings using difflib."""
        return difflib.SequenceMatcher(None, a, b).ratio()

    def evaluate(self, query: str, reference_answer: str) -> Any:
        """
        Evaluate the system using fuzzy matching-based Recall@k and BLEU score.
        
        Args:
            query (str): The query to evaluate.
            reference_answer (str): The expected answer.
            
        Returns:
            dict: A dictionary with evaluation metrics and the generated answer.
        """
        try:
            retrieved = self.qa_chain.retriever.get_relevant_documents(query)
            recall = 0
            for doc in retrieved:
                if self._fuzzy_match(reference_answer.lower(), doc.page_content.lower()) > 0.6:
                    recall = 1
                    break
            generated = self.qa_chain.run(query)
            bleu = sentence_bleu(
                [reference_answer.split()],
                generated.split(),
                weights=(0.5, 0.5, 0, 0)
            )
            return {
                "recall@k": recall,
                "bleu": bleu,
                "generated": generated
            }
        except Exception as e:
            logger.error(f"Evaluation error: {str(e)}\n{traceback.format_exc()}")
            return None

# ---------------------------
# Main Streamlit Application
# ---------------------------
def main() -> None:
    st.set_page_config(page_title="Document QA System", page_icon="📚")
    st.title("Multi-Source Document QA System")

    # Initialize session state for persistence across interactions
    if "qa_system" not in st.session_state:
        st.session_state.qa_system = None
    if "doc_text" not in st.session_state:
        st.session_state.doc_text = None
    
    processor = DocumentProcessor()

    # Sidebar configuration
    with st.sidebar:
        st.header("Configuration")
        max_file_size = st.number_input("Max file size (MB)", min_value=1, value=CONFIG["max_file_size_mb"])
        eval_mode = st.checkbox("Enable Evaluation Mode")
    
    uploaded_file = st.file_uploader(
        "Upload document (PDF, DOCX, TXT)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=False
    )
    
    if uploaded_file:
        if uploaded_file.size > max_file_size * 1024 * 1024:
            st.error(f"File size exceeds {max_file_size}MB limit. Please upload a smaller file.")
            return
        
        try:
            with st.spinner("Processing document..."):
                file_bytes = uploaded_file.getvalue()
                # Use cached extraction; if document is very large, you might call _handle_large_document explicitly.
                text = cached_extract_text(file_bytes, uploaded_file.type)
                st.session_state.doc_text = text  # Cache text in session state
                
                # Optional: if text size is very large, you may process it with _handle_large_document
                if len(text.encode("utf-8")) > 1_000_000:  # > 1MB
                    chunks = processor._handle_large_document(text)
                else:
                    chunks = processor.chunk_text(text)
                
                qa_system = QASystem()
                qa_system.build_index(chunks)
                asyncio.run(qa_system.initialize_qa_chain())
                st.session_state.qa_system = qa_system  # Cache QA system in session state
            st.success("Document processed successfully!")
            
            if eval_mode:
                st.subheader("System Evaluation")
                query = st.text_input("Evaluation Query")
                reference = st.text_area("Reference Answer")
                
                if st.button("Run Evaluation"):
                    results = st.session_state.qa_system.evaluate(query, reference)
                    if results:
                        st.json({
                            "recall@k": results["recall@k"],
                            "bleu_score": round(results["bleu"], 4),
                            "generated_answer": results["generated"]
                        })
                    else:
                        st.error("Evaluation failed. Check logs for details.")
            else:
                st.subheader("Ask Questions")
                query = st.text_input("Enter your question")
                
                if st.button("Get Answer"):
                    with st.spinner("Generating answer..."):
                        try:
                            answer = st.session_state.qa_system.qa_chain.run(query)
                            st.markdown(f"**Answer:**\n{answer}")
                        except Exception as e:
                            logger.error(f"Error generating answer: {str(e)}\n{traceback.format_exc()}")
                            st.error("Error generating answer. Please try again later.")
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}\n{traceback.format_exc()}")
            st.error(f"Error processing document: {str(e)}")

if __name__ == "__main__":
    main()
